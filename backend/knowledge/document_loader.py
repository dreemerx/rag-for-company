"""
多格式文档加载器
- 支持 PDF、Word、Markdown、TXT
- 保留文档结构信息
"""
import os
from typing import List
from pathlib import Path

from .chunker import Document


class DocumentLoader:
    """多格式文档加载器 - 支持 PDF、Word、Markdown、TXT"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".txt"}

    async def load_file(self, file_path: str) -> List[Document]:
        """加载单个文件"""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext in (".docx", ".doc"):
            return self._load_docx(path)
        elif ext == ".md":
            return self._load_markdown(path)
        elif ext == ".txt":
            return self._load_txt(path)

        return []

    async def load_directory(self, dir_path: str) -> List[Document]:
        """加载目录下所有支持的文件"""
        path = Path(dir_path)
        if not path.is_dir():
            raise NotADirectoryError(f"目录不存在: {dir_path}")

        all_docs = []
        for file_path in path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                docs = await self.load_file(str(file_path))
                all_docs.extend(docs)

        return all_docs

    def _load_pdf(self, path: Path) -> List[Document]:
        """加载 PDF 文件"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            docs = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={
                            "source": str(path),
                            "page": i + 1,
                            "type": "pdf",
                            "filename": path.name,
                            "file_type": ".pdf",
                        }
                    ))
            return docs
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")

    def _load_docx(self, path: Path) -> List[Document]:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            if full_text.strip():
                return [Document(
                    page_content=full_text,
                    metadata={
                        "source": str(path),
                        "type": "docx",
                        "filename": path.name,
                        "file_type": ".docx",
                    }
                )]
            return []
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

    def _load_markdown(self, path: Path) -> List[Document]:
        """加载 Markdown 文件"""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        if content.strip():
            return [Document(
                page_content=content,
                metadata={
                    "source": str(path),
                    "type": "markdown",
                    "filename": path.name,
                    "file_type": ".md",
                }
            )]
        return []

    def _load_txt(self, path: Path) -> List[Document]:
        """加载 TXT 文件"""
        # 尝试不同编码
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        content = None

        for encoding in encodings:
            try:
                with open(path, "r", encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content and content.strip():
            return [Document(
                page_content=content,
                metadata={
                    "source": str(path),
                    "type": "txt",
                    "filename": path.name,
                    "file_type": ".txt",
                }
            )]
        return []
